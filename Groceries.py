from tkinter import *
import tkinter.messagebox
root = Tk()
root.title('Groceries')

def startprog():
    global everythingframe
    global startorderingbtn
    global btnrestart
    
    def startordering():
        def submit():
            listofallitems=['quantityrocksalt','quantityfineslt','quantityturmeric','quantitybrownsugar','quantitywhitesugar',
'uriddhallquantity',
'quantitydhalluridsplit',
'quantitytoordhall',
'quantitysplitdhall',
'quantitysplitblackdhall',
'quantitybbengalgramdhall',
'quantitywhitechanna',
'quantitypeanut',
'quantitymustard',
'quantitypepper',
'quantityfenugreek',
'quantityrava',
'quantityidlyrice',
'quantityponnirice',
'quantitybasmatirice',
'quantityrawrice',
'quantityjeerarice',
'quantitygingellyoil',
'quantitygoldwinner',
'quantitycoconutoil',
'quantitycoldpressoil',
'quantitytamerind',
'quantityroastedgramdhall',
'quantityragiflour',
'quantitymoongdhall',
'quantitygreanbean',
'quantityblackchanna',
'quantitymotchaibean',
'quantitypappadam',
'quantityboost',
'quantityfiltercoffee',
'quantityajax',
'quantityhotwaterbag',
'quantitycarrybag',
'quantitysmallonion',
'quantityghee',
'quantitycashewnut',
'quantitymurukkuflour',
'quantityriceflour',
'quantitybengalgramflour',
'quantitycornflour',
'quantitymaidaflour',
'quantityhimalayashampoo',
'quantitygokulsandal',
'quantitygulabjamun',
'quantityaval',
'quantityjaggery',
'quantitycitronpickle',
'quantityincense',
'quantitysambrani',
'quantityprayeroil',
'quantitypanjuthiri',
'quantitysandalpowder',
'quantityvermicelli',
'quantitylg',
'quantitycummin',
'quantityfennelseed',
'quantitygooddaybiscuit',
'quantitybhelpuri',
'quantitycardamon',
'quantitybayleaf',
'quantitycorienderseed',
'quantityjavvarasi']
            global errorframe
            try:
                errorframe.pack_forget()
            except:
                pass
            
            finallist=[]
            if month.get() =='' or year.get()=='':
                errorframe=Frame(orderingframe)
                errorframe.pack(side=BOTTOM)
                errorlabel=Label(errorframe,text='Error! Enter all details', fg='red')
                errorlabel.pack()
            else:
                monthandyear=(month.get()).upper()+' '+year.get()
                for a in range(1,69):
                    if globals()[listofallitems[a-1]].get() == '' or globals()[listofallitems[a-1]].get() == '0':
                        pass
                    else:
                        finallist.append(((globals()['a'+str(a)+'label']).cget("text"),globals()[listofallitems[a-1]].get()))

                import docx
                global everythingframe
                submitbutton.pack_forget()
                btnrestart.pack_forget()
  
                # Create an instance of a word document
                doc = docx.Document()
                doc.add_paragraph('Home Address')

                # Creating a table object
                table = doc.add_table(rows=1, cols=3)

                # Adding heading in the 1st row of the table
                row = table.rows[0].cells
                row[0].text = 'S.NO'
                row[1].text = 'DESCRIPTION'
                row[2].text = 'QUANTITY'
                count=0
                # Adding data from the list to the table
                for item, qty in finallist:

                    # Adding a row and then adding data in it.
                    row = table.add_row().cells
                    # Converting id to string as table can only take string input
                    row[0].text = str(count)
                    row[1].text = item
                    row[2].text = qty
                    count+=1

                # Now save the document to a location
                doc.save('/Users/subramanianponharish/Desktop/Groceries/WordDocs_Program_output/{}.docx'.format(monthandyear))
                everythingframe.pack_forget()
                
                everythingframe=Frame(root)
                everythingframe.pack()


                ##### HEADING IN MAIN PAGE ##############################################################
                headerlabel=Label(everythingframe, text='Groceries', fg='blue',font= "applecherry 35" ) #
                headerlabel.pack()                                                                      #
                spacelbl=Label(everythingframe, text='\n\nORDER SUCCESSFUL\n\n', fg='green',font= "applecherry 20")                                            #
                spacelbl.pack()   
                aspacelbl=Label(everythingframe, text='Doc file is saved in WordDocs_Program output Folder\nFormat it and move it to FINAL DOCUMENTS Folder\n\n', fg='green',font= "applecherry 15")                                            #
                aspacelbl.pack()   
                
                quitbtn=Button(everythingframe, text='Quit',command=root.destroy)

                quitbtn.pack()


            
        global orderingframe
        global quantityrocksalt
        global quantityfineslt
        global quantityturmeric
        global quantitybrownsugar
        global quantitywhitesugar
        global uriddhallquantity
        global quantitydhalluridsplit
        global quantitytoordhall
        global quantitysplitdhall
        global quantitysplitblackdhall
        global quantitybbengalgramdhall
        global quantitywhitechanna
        global quantitypeanut
        global quantitymustard
        global quantitypepper
        global quantityfenugreek
        global quantityrava
        global quantityidlyrice
        global quantityponnirice
        global quantitybasmatirice
        global quantityrawrice
        global quantityjeerarice
        global quantitygingellyoil
        global quantitygoldwinner
        global quantitycoconutoil
        global quantitycoldpressoil
        global quantitytamerind
        global quantityroastedgramdhall
        global quantityragiflour
        global quantitymoongdhall
        global quantitygreanbean
        global quantityblackchanna
        global quantitymotchaibean
        global quantitypappadam
        global quantityboost
        global quantityfiltercoffee
        global quantityajax
        global quantityhotwaterbag
        global quantitycarrybag
        global quantitysmallonion
        global quantityghee
        global quantitycashewnut
        global quantitymurukkuflour
        global quantityriceflour
        global quantitybengalgramflour
        global quantitycornflour
        global quantitymaidaflour
        global quantityhimalayashampoo
        global quantitygokulsandal
        global quantitygulabjamun
        global quantityaval
        global quantityjaggery
        global quantitycitronpickle
        global quantityincense
        global quantitysambrani
        global quantityprayeroil
        global quantitypanjuthiri
        global quantitysandalpowder
        global quantityvermicelli
        global quantitylg
        global quantitycummin
        global quantityfennelseed
        global quantitygooddaybiscuit
        global quantitybhelpuri
        global quantitycardamon
        global quantitybayleaf
        global quantitycorienderseed
        global quantityjavvarasi
        global a1label
        global a2label
        global a3label
        global a4label
        global a5label
        global a6label
        global a7label
        global a8label
        global a9label
        global a10label
        global a11label
        global a12label
        global a13label
        global a14label
        global a15label
        global a16label
        global a17label
        global a18label
        global a19label
        global a20label
        global a21label
        global a22label
        global a23label
        global a24label
        global a25label
        global a26label
        global a27label
        global a28label
        global a29label
        global a30label
        global a31label
        global a32label
        global a33label
        global a34label
        global a35label
        global a36label
        global a37label
        global a38label
        global a39label
        global a40label
        global a41label
        global a42label
        global a43label
        global a44label
        global a45label
        global a46label
        global a47label
        global a48label
        global a49label
        global a50label
        global a51label
        global a52label
        global a53label
        global a54label
        global a55label
        global a56label
        global a57label
        global a58label
        global a59label
        global a60label
        global a61label
        global a62label
        global a63label
        global a64label
        global a65label
        global a66label
        global a67label
        global a68label
        
        startorderingbtn.pack_forget()

        
        orderingframe=Frame(everythingframe)
        orderingframe.pack(side = LEFT)
        
        ##### Month and Year Information #####
        currenttimedetailsframe=Frame(orderingframe)
        currenttimedetailsframe.pack()
        monthlabel=Label(currenttimedetailsframe, text= ' Month:')
        month=Entry(currenttimedetailsframe,width=10) 
        yearlbl=Label(currenttimedetailsframe,text='  Year:')
        year=Entry(currenttimedetailsframe,width=5)
        monthlabel.pack(side= LEFT)
        month.pack(side= LEFT)
        yearlbl.pack(side= LEFT)
        year.pack(side= LEFT)
        ######################################
        spacelabel=Label(orderingframe, text='\n\n')
        spacelabel.pack()
        canvas = Canvas(orderingframe, width=600, height=400)
        
        frameforscroll = Frame(canvas)
        frameforscroll.pack()
        scroll_y = Scrollbar(everythingframe, orient="vertical", command=canvas.yview)
        #####Pack all other frames in frame for scroll
        ##############################################
        #############ITEMS############################
        ##############################################
        #
        rocksaltframe=Frame(frameforscroll)
        rocksaltframe.pack()
        lblframe=Frame(rocksaltframe)
        lblframe.pack()
        
        a1label=Label(lblframe, text='SALT ROCK - 1KG',fg='blue',font= "applecherry 20" , justify=LEFT)
        a1label.pack(side=LEFT)
        qtyframe=Frame(rocksaltframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ')
        quatitylbl.pack(side=LEFT)
        quantityrocksalt=Entry(qtyframe,width=3)
        quantityrocksalt.pack()
        
        spacelabel=Label(frameforscroll, text='----------------------------------------------------------------------\n')
        spacelabel.pack()
        
        #
        finesaltframe=Frame(frameforscroll) #
        finesaltframe.pack() #
        lblframe=Frame(finesaltframe)
        lblframe.pack()
        
        a2label=Label(lblframe, text='FINE SALT - 1KG',fg='blue',font= "applecherry 20" , justify=LEFT) ##
        a2label.pack(side=LEFT)
        qtyframe=Frame(finesaltframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') #
        quatitylbl.pack(side=LEFT)
        quantityfineslt=Entry(qtyframe,width=3) ##
        quantityfineslt.pack() #
        aspacelabel=Label(frameforscroll, text='----------------------------------------------------------------------\n')
        aspacelabel.pack()
        
        ########################################################################################
        turmericframe=Frame(frameforscroll) #
        turmericframe.pack() #
        lblframe=Frame(turmericframe)
        lblframe.pack()
        
        a3label=Label(lblframe, text='TURMERIC',fg='blue',font= "applecherry 20" , justify=LEFT) ##
        a3label.pack(side=LEFT)
        #turmericphoto=PhotoImage(file='/Users/subramanianponharish/Desktop/Groceries/Images/Turmeric.png') ##
        #turmericph=Label(turmericframe,image=turmericphoto) ###
        #turmericph.pack(side=RIGHT) #
        qtyframe=Frame(turmericframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') #
        quatitylbl.pack(side=LEFT)
        quantityturmeric=Entry(qtyframe,width=5) ##
        quantityturmeric.pack() #
        bspacelabel=Label(frameforscroll, text='----------------------------------------------------------------------\n') #
        bspacelabel.pack()
        
        #
        brownsugarframe=Frame(frameforscroll) #
        brownsugarframe.pack() #
        lblframe=Frame(brownsugarframe)
        lblframe.pack()
        
        a4label=Label(lblframe, text='KAFERA BROWN SUGAR – 1KG',fg='blue',font= "applecherry 20" , justify=LEFT) ##
        a4label.pack(side=LEFT)
        qtyframe=Frame(brownsugarframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') #
        quatitylbl.pack(side=LEFT)
        quantitybrownsugar=Entry(qtyframe,width=3) ##
        quantitybrownsugar.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        whitesugarframe=Frame(frameforscroll)  
        whitesugarframe.pack() 
        lblframe=Frame(whitesugarframe)
        lblframe.pack()
        a5label=Label(lblframe, text='SUGAR FINE – 1KG',fg='blue',font= "applecherry 20", justify=LEFT ) 
        a5label.pack(side=LEFT)
        qtyframe=Frame(whitesugarframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitywhitesugar=Entry(qtyframe,width=3) 
        quantitywhitesugar.pack() 
        spacelabel=Label(frameforscroll, text='----------------------------------------------------------------------\n') 
        spacelabel.pack()
        
        #
        uriddhallframe=Frame(frameforscroll) 
        uriddhallframe.pack() 
        lblframe=Frame(uriddhallframe)
        lblframe.pack()
        a6label=Label(lblframe, text='DHALL - URID WHOLE UDHAYAM – 1KG',fg='blue',font= "applecherry 20", justify=LEFT ) #
        a6label.pack(side=LEFT)
        qtyframe=Frame(uriddhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        uriddhallquantity=Entry(qtyframe,width=3) 
        uriddhallquantity.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        dhalluridsplitframe=Frame(frameforscroll) 
        dhalluridsplitframe.pack() 
        lblframe=Frame(dhalluridsplitframe)
        lblframe.pack()
        a7label=Label(lblframe, text='DHALL - URID SPLIT',fg='blue',font= "applecherry 20" , justify=LEFT) 
        a7label.pack(side=LEFT)
        qtyframe=Frame(dhalluridsplitframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitydhalluridsplit=Entry(qtyframe,width=5) 
        quantitydhalluridsplit.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        toordhallframe=Frame(frameforscroll) 
        toordhallframe.pack() 
        lblframe=Frame(toordhallframe)
        lblframe.pack()
        a8label=Label(lblframe, text='DHALL - TOOR UDHAYAM – 1KG',fg='blue',font= "applecherry 20" , justify=LEFT) 
        a8label.pack(side=LEFT)
        qtyframe=Frame(toordhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitytoordhall=Entry(qtyframe,width=3) 
        quantitytoordhall.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        splitdhallframe=Frame(frameforscroll) 
        splitdhallframe.pack() 
        lblframe=Frame(splitdhallframe)
        lblframe.pack()
        a9label=Label(lblframe, text='SPLIT DHALL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a9label.pack(side=LEFT)
        qtyframe=Frame(splitdhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitysplitdhall=Entry(qtyframe,width=5) 
        quantitysplitdhall.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        splitblackdhallframe=Frame(frameforscroll) 
        splitblackdhallframe.pack() 
        lblframe=Frame(splitblackdhallframe)
        lblframe.pack()
        a10label=Label(lblframe, text='SPLIT BLACK DHALL',fg='blue',font= "applecherry 20" , justify=LEFT) 
        a10label.pack(side=LEFT)
        qtyframe=Frame(splitblackdhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitysplitblackdhall=Entry(qtyframe,width=5) 
        quantitysplitblackdhall.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        #
        bengalgramdhallframe=Frame(frameforscroll) 
        bengalgramdhallframe.pack()   
        lblframe=Frame(bengalgramdhallframe)
        lblframe.pack()
        a11label=Label(lblframe, text='BENGAL GRAM DHALL',fg='blue',font= "applecherry 20" , justify=LEFT) 
        a11label.pack(side=LEFT)
        qtyframe=Frame(bengalgramdhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitybbengalgramdhall=Entry(qtyframe,width=5) 
        quantitybbengalgramdhall.pack() 
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        whitechannaframe=Frame(frameforscroll) #
        whitechannaframe.pack() #
        lblframe=Frame(whitechannaframe)#
        lblframe.pack()
        a12label=Label(lblframe, text='WHITE CHANNA',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a12label.pack(side=LEFT)
        qtyframe=Frame(whitechannaframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitywhitechanna=Entry(qtyframe,width=5) #
        quantitywhitechanna.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        peanutframe=Frame(frameforscroll) #
        peanutframe.pack() #
        lblframe=Frame(peanutframe)#
        lblframe.pack()
        a13label=Label(lblframe, text='PEANUT',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a13label.pack(side=LEFT)
        qtyframe=Frame(peanutframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitypeanut=Entry(qtyframe,width=5) #
        quantitypeanut.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        mustardframe=Frame(frameforscroll) #
        mustardframe.pack() #
        lblframe=Frame(mustardframe)#
        lblframe.pack()
        a14label=Label(lblframe, text='MUSTARD',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a14label.pack(side=LEFT)
        qtyframe=Frame(mustardframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitymustard=Entry(qtyframe,width=5) #
        quantitymustard.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        pepperframe=Frame(frameforscroll) #
        pepperframe.pack() #
        lblframe=Frame(pepperframe)#
        lblframe.pack()
        a15label=Label(lblframe, text='PEPPER',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a15label.pack(side=LEFT)
        qtyframe=Frame(pepperframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitypepper=Entry(qtyframe,width=5) #
        quantitypepper.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        fenugreekframe=Frame(frameforscroll) #
        fenugreekframe.pack() #
        lblframe=Frame(fenugreekframe)#
        lblframe.pack()
        a16label=Label(lblframe, text='SELVI FENUGREEK - 250 g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a16label.pack(side=LEFT)
        qtyframe=Frame(fenugreekframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityfenugreek=Entry(qtyframe,width=3) #
        quantityfenugreek.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()

        ravaframe=Frame(frameforscroll) #
        ravaframe.pack() #
        lblframe=Frame(ravaframe)#
        lblframe.pack()
        a17label=Label(lblframe, text='RAVA',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a17label.pack(side=LEFT)
        qtyframe=Frame(ravaframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityrava=Entry(qtyframe,width=5) #
        quantityrava.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        idlyriceframe=Frame(frameforscroll) #
        idlyriceframe.pack() #
        lblframe=Frame(idlyriceframe)#
        lblframe.pack()
        a18label=Label(lblframe, text='SELVI IDLI PONNI - 5KG',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a18label.pack(side=LEFT)
        qtyframe=Frame(idlyriceframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityidlyrice=Entry(qtyframe,width=3) #
        quantityidlyrice.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        ponniriceframe=Frame(frameforscroll) #
        ponniriceframe.pack() #
        lblframe=Frame(ponniriceframe)#
        lblframe.pack()
        a19label=Label(lblframe, text='SELVI PONNI RICE – 5KG ',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a19label.pack(side=LEFT)
        qtyframe=Frame(ponniriceframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityponnirice=Entry(qtyframe,width=3) #
        quantityponnirice.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        basmatiriceframe=Frame(frameforscroll) #
        basmatiriceframe.pack() #
        lblframe=Frame(basmatiriceframe)#
        lblframe.pack()
        a20label=Label(lblframe, text='INDIA GATE BASMATI RICE – 1KG',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a20label.pack(side=LEFT)
        qtyframe=Frame(basmatiriceframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitybasmatirice=Entry(qtyframe,width=3) #
        quantitybasmatirice.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        rawriceframe=Frame(frameforscroll) #
        rawriceframe.pack() #
        lblframe=Frame(rawriceframe)#
        lblframe.pack()
        a21label=Label(lblframe, text='RAW RICE',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a21label.pack(side=LEFT)
        qtyframe=Frame(rawriceframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityrawrice=Entry(qtyframe,width=5) #
        quantityrawrice.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        jeerariceframe=Frame(frameforscroll) #
        jeerariceframe.pack() #
        lblframe=Frame(jeerariceframe)#
        lblframe.pack()
        a22label=Label(lblframe, text='JEERA RICE',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a22label.pack(side=LEFT)
        qtyframe=Frame(jeerariceframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityjeerarice=Entry(qtyframe,width=5) #
        quantityjeerarice.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        gingellyoilframe=Frame(frameforscroll) #
        gingellyoilframe.pack() #
        lblframe=Frame(gingellyoilframe)#
        lblframe.pack()
        a23label=Label(lblframe, text='GINGELLY OIL ',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a23label.pack(side=LEFT)
        qtyframe=Frame(gingellyoilframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygingellyoil=Entry(qtyframe,width=5) #
        quantitygingellyoil.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        goldwinnerframe=Frame(frameforscroll) #
        goldwinnerframe.pack() #
        lblframe=Frame(goldwinnerframe)#
        lblframe.pack()
        a24label=Label(lblframe, text='OIL GOLDWINNER',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a24label.pack(side=LEFT)
        qtyframe=Frame(goldwinnerframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygoldwinner=Entry(qtyframe,width=5) #
        quantitygoldwinner.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        coconutoilframe=Frame(frameforscroll) #
        coconutoilframe.pack() #
        lblframe=Frame(coconutoilframe)#
        lblframe.pack()
        a25label=Label(lblframe, text='PARACHUTE COCONUT OIL – 500g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a25label.pack(side=LEFT)
        qtyframe=Frame(coconutoilframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycoconutoil=Entry(qtyframe,width=3) #
        quantitycoconutoil.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        coldpressoilframe=Frame(frameforscroll) #
        coldpressoilframe.pack() #
        lblframe=Frame(coldpressoilframe)#
        lblframe.pack()
        a26label=Label(lblframe, text='COLD PRESS OIL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a26label.pack(side=LEFT)
        qtyframe=Frame(coldpressoilframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycoldpressoil=Entry(qtyframe,width=5) #
        quantitycoldpressoil.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        tamerindframe=Frame(frameforscroll) #
        tamerindframe.pack() #
        lblframe=Frame(tamerindframe)#
        lblframe.pack()
        a27label=Label(lblframe, text='PS Tamarind',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a27label.pack(side=LEFT)
        qtyframe=Frame(tamerindframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitytamerind=Entry(qtyframe,width=3) #
        quantitytamerind.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        roastedgramdhallframe=Frame(frameforscroll) #
        roastedgramdhallframe.pack() #
        lblframe=Frame(roastedgramdhallframe)#
        lblframe.pack()
        a28label=Label(lblframe, text='ROASTED GRAM DHALL – 500g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a28label.pack(side=LEFT)
        qtyframe=Frame(roastedgramdhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityroastedgramdhall=Entry(qtyframe,width=3) #
        quantityroastedgramdhall.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        ragiflourframe=Frame(frameforscroll) #
        ragiflourframe.pack() #
        lblframe=Frame(ragiflourframe)#
        lblframe.pack()
        a29label=Label(lblframe, text='RAGI FLOUR ',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a29label.pack(side=LEFT)
        qtyframe=Frame(ragiflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityragiflour=Entry(qtyframe,width=3) #
        quantityragiflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        moongdhallframe=Frame(frameforscroll) #
        moongdhallframe.pack() #
        lblframe=Frame(moongdhallframe)#
        lblframe.pack()
        a30label=Label(lblframe, text='MOONG DHALL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a30label.pack(side=LEFT)
        qtyframe=Frame(moongdhallframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitymoongdhall=Entry(qtyframe,width=5) #
        quantitymoongdhall.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        greanbeanframe=Frame(frameforscroll) #
        greanbeanframe.pack() #
        lblframe=Frame(greanbeanframe)#
        lblframe.pack()
        a31label=Label(lblframe, text='GREEN BEAN',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a31label.pack(side=LEFT)
        qtyframe=Frame(greanbeanframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygreanbean=Entry(qtyframe,width=5) #
        quantitygreanbean.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        blackchannaframe=Frame(frameforscroll) #
        blackchannaframe.pack() #
        lblframe=Frame(blackchannaframe)#
        lblframe.pack()
        a32label=Label(lblframe, text='BLACK CHANNA',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a32label.pack(side=LEFT)
        qtyframe=Frame(blackchannaframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityblackchanna=Entry(qtyframe,width=5) #
        quantityblackchanna.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        motchaibeanframe=Frame(frameforscroll) #
        motchaibeanframe.pack() #
        lblframe=Frame(motchaibeanframe)#
        lblframe.pack()
        a33label=Label(lblframe, text='MOTCHAI BEAN',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a33label.pack(side=LEFT)
        qtyframe=Frame(motchaibeanframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitymotchaibean=Entry(qtyframe,width=5) #
        quantitymotchaibean.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        pappadamframe=Frame(frameforscroll) #
        pappadamframe.pack() #
        lblframe=Frame(pappadamframe)#
        lblframe.pack()
        a34label=Label(lblframe, text='TAJ MAHAL PAPPAD',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a34label.pack(side=LEFT)
        qtyframe=Frame(pappadamframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitypappadam=Entry(qtyframe,width=3) #
        quantitypappadam.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        boostframe=Frame(frameforscroll) #
        boostframe.pack() #
        lblframe=Frame(boostframe)#
        lblframe.pack()
        a35label=Label(lblframe, text='BOOST (REFILL)',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a35label.pack(side=LEFT)
        qtyframe=Frame(boostframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityboost=Entry(qtyframe,width=3) #
        quantityboost.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        filtercoffeeframe=Frame(frameforscroll) #
        filtercoffeeframe.pack() #
        lblframe=Frame(filtercoffeeframe)#
        lblframe.pack()
        a36label=Label(lblframe, text='NARASUS UDHAYAM FILTER COFFEE',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a36label.pack(side=LEFT)
        qtyframe=Frame(filtercoffeeframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityfiltercoffee=Entry(qtyframe,width=3) #
        quantityfiltercoffee.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        ajaxframe=Frame(frameforscroll) #
        ajaxframe.pack() #
        lblframe=Frame(ajaxframe)#
        lblframe.pack()
        a37label=Label(lblframe, text='AJAX FLOOR CLEANER',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a37label.pack(side=LEFT)
        qtyframe=Frame(ajaxframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityajax=Entry(qtyframe,width=3) #
        quantityajax.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        hotwaterbagframe=Frame(frameforscroll) #
        hotwaterbagframe.pack() #
        lblframe=Frame(hotwaterbagframe)#
        lblframe.pack()
        a38label=Label(lblframe, text='HOT WATER BAGS (5 X 8)',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a38label.pack(side=LEFT)
        qtyframe=Frame(hotwaterbagframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityhotwaterbag=Entry(qtyframe,width=3) #
        quantityhotwaterbag.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        carrybagframe=Frame(frameforscroll) #
        carrybagframe.pack() #
        lblframe=Frame(carrybagframe)#
        lblframe.pack()
        a39label=Label(lblframe, text='CARRY BAG NO.3',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a39label.pack(side=LEFT)
        qtyframe=Frame(carrybagframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycarrybag=Entry(qtyframe,width=3) #
        quantitycarrybag.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        smallonionframe=Frame(frameforscroll) #
        smallonionframe.pack() #
        lblframe=Frame(smallonionframe)#
        lblframe.pack()
        a40label=Label(lblframe, text='SMALL ONION ',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a40label.pack(side=LEFT)
        qtyframe=Frame(smallonionframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitysmallonion=Entry(qtyframe,width=5) #
        quantitysmallonion.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        gheeframe=Frame(frameforscroll) #
        gheeframe.pack() #
        lblframe=Frame(gheeframe)#
        lblframe.pack()
        a41label=Label(lblframe, text='UDHAYAM GHEE',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a41label.pack(side=LEFT)
        qtyframe=Frame(gheeframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityghee=Entry(qtyframe,width=5) #
        quantityghee.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        cashewnutframe=Frame(frameforscroll) #
        cashewnutframe.pack() #
        lblframe=Frame(cashewnutframe)#
        lblframe.pack()
        a42label=Label(lblframe, text='CASHEW NUT - 250g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a42label.pack(side=LEFT)
        qtyframe=Frame(cashewnutframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycashewnut=Entry(qtyframe,width=3) #
        quantitycashewnut.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        murukkuflourframe=Frame(frameforscroll) #
        murukkuflourframe.pack() #
        lblframe=Frame(murukkuflourframe)#
        lblframe.pack()
        a43label=Label(lblframe, text='LINGAM MURUKKU FLOUR',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a43label.pack(side=LEFT)
        qtyframe=Frame(murukkuflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitymurukkuflour=Entry(qtyframe,width=5) #
        quantitymurukkuflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        riceflourframe=Frame(frameforscroll) #
        riceflourframe.pack() #
        lblframe=Frame(riceflourframe)#
        lblframe.pack()
        a44label=Label(lblframe, text='RICE FLOUR',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a44label.pack(side=LEFT)
        qtyframe=Frame(riceflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityriceflour=Entry(qtyframe,width=5) #
        quantityriceflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        bengalgramflourframe=Frame(frameforscroll) #
        bengalgramflourframe.pack() #
        lblframe=Frame(bengalgramflourframe)#
        lblframe.pack()
        a45label=Label(lblframe, text='BENGAL GRAM FLOUR',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a45label.pack(side=LEFT)
        qtyframe=Frame(bengalgramflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitybengalgramflour=Entry(qtyframe,width=5) #
        quantitybengalgramflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        cornflourframe=Frame(frameforscroll) #
        cornflourframe.pack() #
        lblframe=Frame(cornflourframe)#
        lblframe.pack()
        a46label=Label(lblframe, text='CORN FLOUR',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a46label.pack(side=LEFT)
        qtyframe=Frame(cornflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycornflour=Entry(qtyframe,width=5) #
        quantitycornflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        maidaflourframe=Frame(frameforscroll) #
        maidaflourframe.pack() #
        lblframe=Frame(maidaflourframe)#
        lblframe.pack()
        a47label=Label(lblframe, text='MAIDA FLOUR',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a47label.pack(side=LEFT)
        qtyframe=Frame(maidaflourframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitymaidaflour=Entry(qtyframe,width=5) #
        quantitymaidaflour.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        himalayashampooframe=Frame(frameforscroll) #
        himalayashampooframe.pack() #
        lblframe=Frame(himalayashampooframe)#
        lblframe.pack()
        a48label=Label(lblframe, text='HIMALAYA SHAMPOO',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a48label.pack(side=LEFT)
        qtyframe=Frame(himalayashampooframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityhimalayashampoo=Entry(qtyframe,width=3) #
        quantityhimalayashampoo.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        gokulsandalframe=Frame(frameforscroll) #
        gokulsandalframe.pack() #
        lblframe=Frame(gokulsandalframe)#
        lblframe.pack()
        a49label=Label(lblframe, text='GOKUL SANDAL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a49label.pack(side=LEFT)
        qtyframe=Frame(gokulsandalframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygokulsandal=Entry(qtyframe,width=3) #
        quantitygokulsandal.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        gulabjamunframe=Frame(frameforscroll) #
        gulabjamunframe.pack() #
        lblframe=Frame(gulabjamunframe)#
        lblframe.pack()
        a50label=Label(lblframe, text='MTR GULAB JAMUN MIX',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a50label.pack(side=LEFT)
        qtyframe=Frame(gulabjamunframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygulabjamun=Entry(qtyframe,width=3) #
        quantitygulabjamun.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        avalframe=Frame(frameforscroll) #
        avalframe.pack() #
        lblframe=Frame(avalframe)#
        lblframe.pack()
        a51label=Label(lblframe, text='AVAL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a51label.pack(side=LEFT)
        qtyframe=Frame(avalframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityaval=Entry(qtyframe,width=5) #
        quantityaval.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        jaggeryframe=Frame(frameforscroll) #
        jaggeryframe.pack() #
        lblframe=Frame(jaggeryframe)#
        lblframe.pack()
        a52label=Label(lblframe, text='JAGGERY – 500g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a52label.pack(side=LEFT)
        qtyframe=Frame(jaggeryframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityjaggery=Entry(qtyframe,width=3) #
        quantityjaggery.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        citronpickleframe=Frame(frameforscroll) #
        citronpickleframe.pack() #
        lblframe=Frame(citronpickleframe)#
        lblframe.pack()
        a53label=Label(lblframe, text='CITRON PICKLE',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a53label.pack(side=LEFT)
        qtyframe=Frame(citronpickleframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycitronpickle=Entry(qtyframe,width=3) #
        quantitycitronpickle.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        incenseframe=Frame(frameforscroll) #
        incenseframe.pack() #
        lblframe=Frame(incenseframe)#
        lblframe.pack()
        a54label=Label(lblframe, text='CYCLE 3 IN 1 - 100 STICKS',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a54label.pack(side=LEFT)
        qtyframe=Frame(incenseframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityincense=Entry(qtyframe,width=3) #
        quantityincense.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        sambraniframe=Frame(frameforscroll) #
        sambraniframe.pack() #
        lblframe=Frame(sambraniframe)#
        lblframe.pack()
        a55label=Label(lblframe, text='AMBAL SAMBRANI',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a55label.pack(side=LEFT)
        qtyframe=Frame(sambraniframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitysambrani=Entry(qtyframe,width=3) #
        quantitysambrani.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        prayeroilframe=Frame(frameforscroll) #
        prayeroilframe.pack() #
        lblframe=Frame(prayeroilframe)#
        lblframe.pack()
        a56label=Label(lblframe, text='PRAYER OIL',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a56label.pack(side=LEFT)
        qtyframe=Frame(prayeroilframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityprayeroil=Entry(qtyframe,width=3) #
        quantityprayeroil.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        panjuthiriframe=Frame(frameforscroll) #
        panjuthiriframe.pack() #
        lblframe=Frame(panjuthiriframe)#
        lblframe.pack()
        a57label=Label(lblframe, text='PANJU THIRI',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a57label.pack(side=LEFT)
        qtyframe=Frame(panjuthiriframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitypanjuthiri=Entry(qtyframe,width=3) #
        quantitypanjuthiri.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        sandalpowderframe=Frame(frameforscroll) #
        sandalpowderframe.pack() #
        lblframe=Frame(sandalpowderframe)#
        lblframe.pack()
        a58label=Label(lblframe, text='SANDAL POWDER',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a58label.pack(side=LEFT)
        qtyframe=Frame(sandalpowderframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitysandalpowder=Entry(qtyframe,width=3) #
        quantitysandalpowder.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        vermicelliframe=Frame(frameforscroll) #
        vermicelliframe.pack() #
        lblframe=Frame(vermicelliframe)#
        lblframe.pack()
        a59label=Label(lblframe, text='BAMBINO VERMICELLI – 350 g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a59label.pack(side=LEFT)
        qtyframe=Frame(vermicelliframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityvermicelli=Entry(qtyframe,width=3) #
        quantityvermicelli.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        lgframe=Frame(frameforscroll) #
        lgframe.pack() #
        lblframe=Frame(lgframe)#
        lblframe.pack()
        a60label=Label(lblframe, text='LG ASAFEOTIDA POWDER – 100 g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a60label.pack(side=LEFT)
        qtyframe=Frame(lgframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitylg=Entry(qtyframe,width=3) #
        quantitylg.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        cumminframe=Frame(frameforscroll) #
        cumminframe.pack() #
        lblframe=Frame(cumminframe)#
        lblframe.pack()
        a61label=Label(lblframe, text='CUMIN SEED',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a61label.pack(side=LEFT)
        qtyframe=Frame(cumminframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycummin=Entry(qtyframe,width=3) #
        quantitycummin.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        fennelseedframe=Frame(frameforscroll) #
        fennelseedframe.pack() #
        lblframe=Frame(fennelseedframe)#
        lblframe.pack()
        a62label=Label(lblframe, text='FENNEL SEED',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a62label.pack(side=LEFT)
        qtyframe=Frame(fennelseedframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityfennelseed=Entry(qtyframe,width=3) #
        quantityfennelseed.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        gooddaybiscuitframe=Frame(frameforscroll) #
        gooddaybiscuitframe.pack() #
        lblframe=Frame(gooddaybiscuitframe)#
        lblframe.pack()
        a63label=Label(lblframe, text='GOOD DAY BISCUIT',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a63label.pack(side=LEFT)
        qtyframe=Frame(gooddaybiscuitframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitygooddaybiscuit=Entry(qtyframe,width=3) #
        quantitygooddaybiscuit.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        bhelpuriframe=Frame(frameforscroll) #
        bhelpuriframe.pack() #
        lblframe=Frame(bhelpuriframe)#
        lblframe.pack()
        a64label=Label(lblframe, text='BHEL PURI',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a64label.pack(side=LEFT)
        qtyframe=Frame(bhelpuriframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitybhelpuri=Entry(qtyframe,width=3) #
        quantitybhelpuri.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        cardamonframe=Frame(frameforscroll) #
        cardamonframe.pack() #
        lblframe=Frame(cardamonframe)#
        lblframe.pack()
        a65label=Label(lblframe, text='CARADAMON SEED',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a65label.pack(side=LEFT)
        qtyframe=Frame(cardamonframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycardamon=Entry(qtyframe,width=3) #
        quantitycardamon.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        bayleafframe=Frame(frameforscroll) #
        bayleafframe.pack() #
        lblframe=Frame(bayleafframe)#
        lblframe.pack()
        a66label=Label(lblframe, text='BAY LEAF',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a66label.pack(side=LEFT)
        qtyframe=Frame(bayleafframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitybayleaf=Entry(qtyframe,width=3) #
        quantitybayleaf.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        corienderseedframe=Frame(frameforscroll) #
        corienderseedframe.pack() #
        lblframe=Frame(corienderseedframe)#
        lblframe.pack()
        a67label=Label(lblframe, text='CORIENDER SEED',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a67label.pack(side=LEFT)
        qtyframe=Frame(corienderseedframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantitycorienderseed=Entry(qtyframe,width=3) #
        quantitycorienderseed.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        javvarasiframe=Frame(frameforscroll) #
        javvarasiframe.pack() #
        lblframe=Frame(javvarasiframe)#
        lblframe.pack()
        a68label=Label(lblframe, text='BIG JAVVARAISI – 250g',fg='blue',font= "applecherry 20" , justify=LEFT) #
        a68label.pack(side=LEFT)
        qtyframe=Frame(javvarasiframe)#
        qtyframe.pack()
        quatitylbl=Label(qtyframe, text='Quantity: ') 
        quatitylbl.pack(side=LEFT)
        quantityjavvarasi=Entry(qtyframe,width=3) #
        quantityjavvarasi.pack() #
        spacelabel=Label(frameforscroll, text='-----------------------------------------------------------------------\n') #
        spacelabel.pack()
        
        
        
        
        # make sure everything is displayed before configuring the scrollregion
        canvas.create_window(0, 0, anchor='nw', window=frameforscroll)
        canvas.update_idletasks()

        canvas.configure(scrollregion=canvas.bbox('all'), yscrollcommand=scroll_y.set)
                 
        canvas.pack(fill='both', expand=True)
        scroll_y.pack(fill='y', side='right')
        
        submitbutton=Button(orderingframe, text='Submit', command=submit)
        submitbutton.pack()
        
        root.mainloop()
        

    
    
    def restartprog():
        ans=tkinter.messagebox.askquestion('GROCERIES','Are you sure you want to reset?')
        if ans == 'yes':
            everythingframe.pack_forget()
            startprog()
    
    try:
        everythingframe.pack_forget()
    except:
        pass
    everythingframe=Frame(root)
    everythingframe.pack()
    
    
    ##### HEADING IN MAIN PAGE ##############################################################
    headerlabel=Label(everythingframe, text='Groceries', fg='blue',font= "applecherry 35" ) #
    headerlabel.pack()                                                                      #
    spacelbl=Label(everythingframe, text='\n\n')                                            #
    spacelbl.pack()                                                                         #
    #########################################################################################
    ##### BUTTONS IN MAIN PAGE #####
    startorderingbtn=Button(everythingframe,text='Start Ordering',command=startordering)
    startorderingbtn.pack()

    btnrestart=Button(everythingframe,text='restart',command=restartprog)
    btnrestart.pack(side=BOTTOM)
    root.mainloop()
    #################################
    

startprog()



####
#print(label.cget("text")) --> use it to get text from label
